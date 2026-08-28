"""
Renders structured resume content (or plain cover-letter text) to PDF or
DOCX bytes. Pure formatting — no AI calls, no content generation. Kept
deliberately simple (no template engine) so the output is predictable and
easy to keep ATS-friendly (single column, standard fonts, no tables/text
boxes that ATS parsers choke on).
"""
import io

from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.domain.schemas.resume import ResumeContent


class DocumentExportService:
    # --- Resume export ---

    @staticmethod
    def resume_to_pdf(content: ResumeContent) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=LETTER, topMargin=0.6 * inch, bottomMargin=0.6 * inch,
            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        )
        styles = getSampleStyleSheet()
        name_style = ParagraphStyle("Name", parent=styles["Title"], fontSize=18, spaceAfter=4)
        contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=9, textColor="#444444")
        heading_style = ParagraphStyle(
            "SectionHeading", parent=styles["Heading2"], fontSize=12, spaceBefore=12, spaceAfter=4,
        )
        body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=13)
        bullet_style = ParagraphStyle("Bullet", parent=body_style, leftIndent=14, bulletIndent=4)

        elements = [Paragraph(content.contact.full_name, name_style)]

        contact_parts = [
            p for p in [
                content.contact.email, content.contact.phone, content.contact.location,
                content.contact.linkedin_url, content.contact.github_url, content.contact.portfolio_url,
            ] if p
        ]
        if contact_parts:
            elements.append(Paragraph(" · ".join(contact_parts), contact_style))

        if content.summary:
            elements.append(Paragraph("Summary", heading_style))
            elements.append(Paragraph(content.summary, body_style))

        if content.skills:
            elements.append(Paragraph("Skills", heading_style))
            elements.append(Paragraph(", ".join(content.skills), body_style))

        if content.experience:
            elements.append(Paragraph("Experience", heading_style))
            for exp in content.experience:
                date_range = " — ".join(p for p in [exp.start_date, exp.end_date] if p)
                header = f"<b>{exp.title}</b>, {exp.company}"
                if date_range:
                    header += f" &nbsp;&nbsp;<font color='#666666'>{date_range}</font>"
                elements.append(Paragraph(header, body_style))
                for bullet in exp.bullets:
                    elements.append(Paragraph(f"• {bullet}", bullet_style))
                elements.append(Spacer(1, 6))

        if content.projects:
            elements.append(Paragraph("Projects", heading_style))
            for proj in content.projects:
                title = f"<b>{proj.name}</b>"
                if proj.tech_stack:
                    title += f" &nbsp;<font color='#666666'>({', '.join(proj.tech_stack)})</font>"
                elements.append(Paragraph(title, body_style))
                for bullet in proj.bullets:
                    elements.append(Paragraph(f"• {bullet}", bullet_style))
                elements.append(Spacer(1, 6))

        if content.education:
            elements.append(Paragraph("Education", heading_style))
            for edu in content.education:
                date_range = " — ".join(p for p in [edu.start_date, edu.end_date] if p)
                line = f"<b>{edu.institution}</b>"
                if edu.degree:
                    line += f", {edu.degree}"
                if edu.field_of_study:
                    line += f" in {edu.field_of_study}"
                if date_range:
                    line += f" &nbsp;&nbsp;<font color='#666666'>{date_range}</font>"
                elements.append(Paragraph(line, body_style))

        if content.achievements:
            elements.append(Paragraph("Achievements", heading_style))
            for achievement in content.achievements:
                elements.append(Paragraph(f"• {achievement}", bullet_style))

        doc.build(elements)
        return buffer.getvalue()

    @staticmethod
    def resume_to_docx(content: ResumeContent) -> bytes:
        doc = DocxDocument()

        title = doc.add_heading(content.contact.full_name, level=0)
        title.runs[0].font.size = Pt(20)

        contact_parts = [
            p for p in [
                content.contact.email, content.contact.phone, content.contact.location,
                content.contact.linkedin_url, content.contact.github_url, content.contact.portfolio_url,
            ] if p
        ]
        if contact_parts:
            doc.add_paragraph(" · ".join(contact_parts))

        if content.summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(content.summary)

        if content.skills:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(content.skills))

        if content.experience:
            doc.add_heading("Experience", level=1)
            for exp in content.experience:
                date_range = " — ".join(p for p in [exp.start_date, exp.end_date] if p)
                p = doc.add_paragraph()
                p.add_run(f"{exp.title}, {exp.company}").bold = True
                if date_range:
                    p.add_run(f"  ({date_range})")
                for bullet in exp.bullets:
                    doc.add_paragraph(bullet, style="List Bullet")

        if content.projects:
            doc.add_heading("Projects", level=1)
            for proj in content.projects:
                p = doc.add_paragraph()
                p.add_run(proj.name).bold = True
                if proj.tech_stack:
                    p.add_run(f"  ({', '.join(proj.tech_stack)})")
                for bullet in proj.bullets:
                    doc.add_paragraph(bullet, style="List Bullet")

        if content.education:
            doc.add_heading("Education", level=1)
            for edu in content.education:
                date_range = " — ".join(p for p in [edu.start_date, edu.end_date] if p)
                line = edu.institution
                if edu.degree:
                    line += f", {edu.degree}"
                if edu.field_of_study:
                    line += f" in {edu.field_of_study}"
                if date_range:
                    line += f"  ({date_range})"
                doc.add_paragraph(line)

        if content.achievements:
            doc.add_heading("Achievements", level=1)
            for achievement in content.achievements:
                doc.add_paragraph(achievement, style="List Bullet")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # --- Cover letter export ---

    @staticmethod
    def cover_letter_to_pdf(letter_text: str) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=LETTER, topMargin=1 * inch, bottomMargin=1 * inch)
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=12)

        elements = [Paragraph(para.replace("\n", "<br/>"), body_style) for para in letter_text.split("\n\n") if para.strip()]
        doc.build(elements)
        return buffer.getvalue()

    @staticmethod
    def cover_letter_to_docx(letter_text: str) -> bytes:
        doc = DocxDocument()
        for para in letter_text.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
